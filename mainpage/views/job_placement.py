import csv
from django.urls import reverse 
from django.http import HttpResponse, Http404
from weasyprint import HTML
import mimetypes 
from django.urls import reverse
import zipfile
from django.conf import settings
from django.shortcuts import render, redirect
from django.db.models import Q, Value
from django.db.models.functions import Concat
from django.views import View
from django.db.models import Q
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.decorators import login_required
from django.contrib.auth.backends import ModelBackend
from django.core.exceptions import ObjectDoesNotExist as ObjException
from django.contrib import messages
from datetime import datetime
from django.utils import timezone

from django.views.decorators.csrf import csrf_exempt
from ..models import ( SeminarAttendance, TransactionReport, studentInfo, 
                     JobPlacementAdminUser, Seminar, OJTCompany, OJTStudent,
                     OJTRequirements,
                    )
from ..forms import ( SeminarForm, SeminarAttendanceForm, TransactionForm, 
              AdminSignUpForm,  # StudentSignUpForm, StudentLoginForm,
                    OjtHiringForm, OJTStudentForm, EmailAuthenticationForm, 
                    OJTRequirementsForm, StatusWidget, ScrapperFile
                    )

# Seminar Attendance AJAX post request
import json
from django.http import JsonResponse
from django.urls import reverse
from django.http import HttpResponse
from django.template.loader import render_to_string
from weasyprint import HTML
from django.db import transaction
import os
import time
# autogenerate requirements
from docx import Document
from django.db import models

from io import BytesIO
from zipfile import ZipFile
from django.http import FileResponse, Http404
ojtrequirements_application_letter = './media/templates/ojt_requirements/application_letter.docx'
ojtrequirements_biodata = './media/templates/ojt_requirements/biodata.docx'
ojtrequirements_endorsement_letter = './media/templates/ojt_requirements/endorsement_letter.docx'
ojtrequirements_medical = './media/templates/ojt_requirements/medical_clearance.docx'
ojtrequirements_moa = './media/templates/ojt_requirements/moa.docx'
ojtrequirements_template_output_path = './media/templates/ojt_requirements/generated'
# from mainpage.decorators.decorators import sao_admin_required@login_required 
def stream_ojt_file(request, req_id, attr_name):
    """
    Reads a file from storage and streams it to the browser.
    If it's a PDF/Image, it shows inline.
    If it's a .docx or other, it returns an error page.
    """
    
    # Define file types that browsers can show in a tab
    ALLOWED_INLINE_TYPES = [
        'application/pdf',
        'image/jpeg',
        'image/png',
        'image/gif',
        'image/bmp',
        'image/webp'
    ]
    
    try:
        req = OJTRequirements.objects.get(pk=req_id)

        FIELD_MAP = {
            'nondis': 'non_disclosure',
            'biodata': 'biodata',
            'consent': 'parents_consent',
            'apl_letter': 'application_letter',
            'medical': 'medical',
            'moa': 'moa',
            'endorsement': 'endorsement',
            'cert': 'certification',
        }

        model_field_name = FIELD_MAP.get(attr_name)
        if not model_field_name:
            raise Http404("Invalid file type key")

        file_field = getattr(req, model_field_name, None)

        if file_field and file_field.storage.exists(file_field.name):
            
            # 1. Guess the file's content type
            content_type, _ = mimetypes.guess_type(file_field.name)
            
            if not content_type:
                content_type = 'application/octet-stream' # Default for unknown
            
            # 2. Check if the type is viewable
            if content_type in ALLOWED_INLINE_TYPES:
                # It's a PDF or Image: Stream it to be viewed in the tab
                file_content = file_field.read()
                response = HttpResponse(file_content, content_type=content_type)
                response['Content-Disposition'] = f'inline; filename="{file_field.name}"'
                return response
                
            else:
                # It's NOT a viewable file (like .docx)
                # Return an HTML error message to display in the new tab
                error_html = f"""
                <body style="font-family: sans-serif; padding: 40px;">
                    <h2 style="color: #c0392b;">Cannot Preview File</h2>
                    <p>The uploaded file (<b>{file_field.name}</b>) cannot be previewed in the browser because it is not a PDF or image.</p>
                    <p><b>Detected file type:</b> {content_type}</p>
                    <p>This file will need to be downloaded to be viewed. You can close this tab and 'Decline' the file if it is not in the correct format.</p>
                </body>
                """
                return HttpResponse(error_html, content_type="text/html", status=200)

        else:
            raise Http404("File not found")

    except OJTRequirements.DoesNotExist:
        raise Http404("Requirement record not found")
    except Exception as e:
        print(f"Error streaming file: {e}")
        return HttpResponse(f"<h1>Error processing file</h1><p>{e}</p>", status=500)

# === ADD THIS FUNCTION BACK ===
@csrf_exempt
def view_pdf(request, id):
    """
    This is called by AJAX. It returns a JSON response with the 
    URL to the 'stream_ojt_file' view.
    """
    try:
        req = OJTRequirements.objects.get(ojt_requirement_id = id)
    except OJTRequirements.DoesNotExist:
        error = {'error': "OJT Requirement NOT Found"}
        return JsonResponse(error, status=404)
        
    pdf = None 
    
    if request.method == 'POST':
        attr_name = request.POST.get('attr_name') 
        
        # This mapping is critical. It maps the 'name' from the
        # button to the actual file field on the model.
        FIELD_MAP = {
            'nondis': req.non_disclosure,
            'biodata': req.biodata,
            'consent': req.parents_consent,
            'apl_letter': req.application_letter,
            'medical': req.medical,
            'moa': req.moa,
            'endorsement': req.endorsement,
            'cert': req.certification,
        }

        pdf_field = FIELD_MAP.get(attr_name)

        if not pdf_field:
            error = {'error': "Invalid file type specified"}
            return JsonResponse(error, status=400)

        if pdf_field and pdf_field.url:
            pdf_data = {
                # We send the URL to our 'stream_ojt_file' view
                'url': reverse('stream_ojt_file', args=[req.ojt_requirement_id, attr_name]),
                'name': pdf_field.name,
            }
            return JsonResponse(pdf_data)
        else:
            error = {'error': "File not found or not uploaded"}
            return JsonResponse(error, status=44)
    
    error = {'error': "Invalid request method. Must be POST."}
    return JsonResponse(error, status=405)


# === KEEP THIS FUNCTION ===
@login_required 
def stream_ojt_file(request, req_id, attr_name):
    """
    Reads a file from storage and streams it to the browser.
    If it's a PDF/Image, it shows inline.
    If it's a .docx or other, it returns an error page.
    """
    
    ALLOWED_INLINE_TYPES = [
        'application/pdf',
        'image/jpeg',
        'image/png',
        'image/gif',
        'image/bmp',
        'image/webp'
    ]
    
    try:
        req = OJTRequirements.objects.get(pk=req_id)

        FIELD_MAP = {
            'nondis': 'non_disclosure',
            'biodata': 'biodata',
            'consent': 'parents_consent',
            'apl_letter': 'application_letter',
            'medical': 'medical',
            'moa': 'moa',
            'endorsement': 'endorsement',
            'cert': 'certification',
        }

        model_field_name = FIELD_MAP.get(attr_name)
        if not model_field_name:
            raise Http404("Invalid file type key")

        file_field = getattr(req, model_field_name, None)

        if file_field and file_field.storage.exists(file_field.name):
            
            content_type, _ = mimetypes.guess_type(file_field.name)
            if not content_type:
                content_type = 'application/octet-stream' 
            
            if content_type in ALLOWED_INLINE_TYPES:
                # PDF/Image: Show in iframe
                file_content = file_field.read()
                response = HttpResponse(file_content, content_type=content_type)
                response['Content-Disposition'] = f'inline; filename="{file_field.name}"'
                return response
                
            else:
                # .DOCX or other: Show error in iframe
                error_html = f"""
                <div style="font-family: sans-serif; padding: 20px;">
                    <h2>Cannot Display File</h2>
                    <p>The uploaded file (<b>{file_field.name}</b>) cannot be previewed because it is not a PDF or image.</p>
                    <p><b>Detected file type:</b> {content_type}</p>
                    <hr>
                    <p style="color: #c0392b; font-weight: bold;">
                        Please close this preview and 'Decline' this file.
                    </p>
                </div>
                """
                return HttpResponse(error_html, content_type="text/html", status=200)
        else:
            raise Http404("File not found")

    except OJTRequirements.DoesNotExist:
        raise Http404("Requirement record not found")
    except Exception as e:
        print(f"Error streaming file: {e}")
        return HttpResponse(f"<h1>Error processing file</h1><p>{e}</p>", status=500)
@login_required 
def stream_ojt_file(request, req_id, attr_name):
    """
    Reads a file from storage and streams it to the browser
    with the correct Content-Type to be displayed in an iframe.
    """
    try:
        req = OJTRequirements.objects.get(pk=req_id)

        FIELD_MAP = {
            'nondis': 'non_disclosure',
            'biodata': 'biodata',
            'consent': 'parents_consent',
            'apl_letter': 'application_letter',
            'medical': 'medical',
            'moa': 'moa',
            'endorsement': 'endorsement',
            'cert': 'certification',
        }

        model_field_name = FIELD_MAP.get(attr_name)
        if not model_field_name:
            raise Http44("Invalid file type")

        file_field = getattr(req, model_field_name, None)

        if file_field and file_field.storage.exists(file_field.name):
            # Read the file's content
            file_content = file_field.read()
            content_type, _ = mimetypes.guess_type(file_field.name)
            
        
            if not content_type:
                content_type = 'application/pdf' # A safer default for this app


            response = HttpResponse(file_content, content_type=content_type)
            
            response['Content-Disposition'] = f'inline; filename="{file_field.name}"'
            
            return response
        else:
            raise Http404("File not found")

    except OJTRequirements.DoesNotExist:
        raise Http404("Requirement record not found")
    except Exception as e:
        print(f"Error streaming file: {e}")
        raise Http404("Error processing file")
@login_required
def get_ojt_pdf_url(request, req_id, attr_name):
    """
    Safely gets the URL for a specific file field on an OJTRequirement object.
    """
    if not request.method == "GET":
        return JsonResponse({"error": "Invalid request method"}, status=405)

    try:
        # Get the main requirement record
        req = get_object_or_404(OJTRequirements, pk=req_id)

        # A mapping of the 'name' from the HTML button to the actual model field name
        # This is a security measure to prevent arbitrary attribute access.
        FIELD_MAP = {
            'nondis': 'non_disclosure',
            'biodata': 'biodata',
            'consent': 'parents_consent',
            'apl_letter': 'application_letter',
            'medical': 'medical',
            'moa': 'moa',
            'endorsement': 'endorsement',
            'cert': 'certification',
        }

        # Get the correct model field name from our map
        model_field_name = FIELD_MAP.get(attr_name)

        if not model_field_name:
            return JsonResponse({"error": "Invalid file type specified"}, status=400)

        # Get the file field from the requirement object (e.g., req.non_disclosure)
        file_field = getattr(req, model_field_name, None)

        if file_field and file_field.url:
            # If the file exists, return its URL
            return JsonResponse({"url": file_field.url})
        else:
            # If the field is empty or has no file
            return JsonResponse({"error": "File not found or not uploaded"}, status=404)

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
def get_user_backend(user):
    if isinstance(user, JobPlacementAdminUser):
        return 'jobplacement.auth_backends.AdminUserBackend'
    elif isinstance(user, studentInfo):
        return 'jobplacement.auth_backends.studentInfoBackend'
    return None
#deleteable
from django.contrib.auth.hashers import check_password
import os # <-- Add this import at the top of your file
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
# ... (other imports)
from ..models import studentInfo, OJTCompany, OJTStudent # Adjust as needed
from ..forms import OJTStudentForm # Adjust as needed
# ... (your template path variables)
from django.db.models import Q
from django.core.paginator import Paginator
from urllib.parse import urlencode


def admin_student_tracker(request):
    # Check if user is an admin
    if not (request.user.role != 'student' or request.user.is_superuser or getattr(request.user, 'role', None) == 'guard'):
        messages.error(request, "You do not have permission to view this page.")
        return redirect('admin_login') # Or student home page

    base_template = "adminmain.html"
    
    # Get search query
    search_query = request.GET.get('search', '')
    
    # Start with all assigned students
    assignment_list = OJTStudent.objects.select_related('studID', 'company_id').all()

    # Apply search filter
    if search_query:
        assignment_list = assignment_list.filter(
            Q(studID__studID__icontains=search_query) |
            Q(studID__firstname__icontains=search_query) |
            Q(studID__lastname__icontains=search_query) |
            Q(company_id__company_name__icontains=search_query)
        ).distinct()

    # Apply ordering
    assignment_list = assignment_list.order_by('-date_started') # Show most recent first

    # Pagination
    paginator = Paginator(assignment_list, 15) # 15 assignments per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Prepare query string for pagination links
    query_params = request.GET.copy()
    if 'page' in query_params:
        del query_params['page']
    base_query_string = query_params.urlencode()

    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'base_query_string': base_query_string,
        'base_template': base_template,
    }
    return render(request, 'jobplacement/admin_student_tracker.html', context)

def ojt_requiremets_download(request):
    """
    This page shows a success message and provides the download link 
    for the zip file created by ojt_assign_student.
    """
    # Get the data from the session, using .pop() to clear it
    download_url = request.session.pop('download_url', None)
    student_name = request.session.pop('assigned_student_name', 'The student')
    company_name = request.session.pop('assigned_company_name', 'their company')

    # If the user lands on this page without a file, redirect them
    if not download_url:
        messages.warning(request, "No download file specified.")
        return redirect('ojt_hiring')
    
    base_template = "adminmain.html" # Or your logic to find the base template
    
    context = {
        'download_url': download_url,
        'student_name': student_name,
        'company_name': company_name,
        'base_template': base_template,
    }
    return render(request, 'jobplacement/ojt_download_page.html', context)

def ojt_assign_student(request):
    base_template = "adminmain.html" if (request.user.role != 'student' or request.user.is_superuser or getattr(request.user, 'role', None) == 'guard') else "main.html"    # --- Access control ---
    if not (request.user.role != 'student' or request.user.is_superuser):
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')

    if not (isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser):
        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')

    if request.method == 'POST':
        print("Processing OJT assignment...")
        # --- Get POST data safely ---
        student_id = request.POST.get('studID')
        company_id = request.POST.get('company_id')
        endorser = request.POST.get('endorser_name')
        endorser_num = request.POST.get('endorser_num')
        endorser_email = request.POST.get('endorser_email')
        endorser_program = request.POST.get('endorser_program')
        duration = request.POST.get("duration")
        print(f"Received data: student_id={student_id}, company_id={company_id}, endorser={endorser}, duration={duration}")
        params = {
            'endorser_name': endorser,
            'endorser_num': endorser_num,
            'endorser_email': endorser_email,
            'endorser_program': endorser_program,
        }

        # --- Validate student and company existence ---
        try:
            student = studentInfo.objects.get(studID=student_id)
        except studentInfo.DoesNotExist:
            messages.error(request, f"Student with ID '{student_id}' not found.")
            return redirect('ojt_hiring')
        except ValueError:
             messages.error(request, f"Invalid Student ID format: '{student_id}'.")
             return redirect('ojt_hiring')

        try:
            company = OJTCompany.objects.get(company_id=company_id)
        except OJTCompany.DoesNotExist:
            messages.error(request, f"Company with ID '{company_id}' not found.")
            return redirect('ojt_hiring')
        except ValueError:
             messages.error(request, f"Invalid Company ID format: '{company_id}'.")
             return redirect('ojt_hiring')

        # --- Check slots ---
        # Note: Your model's .save() method handles this, but this is a good pre-check.
        if company.number_of_slots < 1:
            messages.error(request, f"No slots left for {company.company_name}.")
            return redirect('ojt_hiring')

        # --- Create OJT assignment ---
        try:
            new_ojt_form = OJTStudentForm(request.POST) # Pass request.POST here
            
            if new_ojt_form.is_valid():
                new_ojt = new_ojt_form.save(commit=False)
                new_ojt.studID = student
                new_ojt.company_id = company
                new_ojt.save() # This will trigger your .save() method and subtract the slot

                messages.success(request, f"Student {student.firstname} {student.lastname} assigned successfully to {company.company_name}.")
                # log_activity(user=request.user, action=f"Assigned {student.firstname} {student.lastname} to {company.company_name}") # Uncomment if log_activity is defined

                # --- Generate documents with Enhanced Error Trapping ---
                generation_successful = True # Flag to track if all documents generated
                try:
                    print(f"Attempting to generate documents in: {ojtrequirements_template_output_path}")
                    os.makedirs(ojtrequirements_template_output_path, exist_ok=True) 

                    print("Generating Application Letter...")
                    gen_application_letter(ojtrequirements_application_letter, ojtrequirements_template_output_path, student_id, company_id)

                    print("Generating Biodata...")
                    gen_biodata(ojtrequirements_biodata, ojtrequirements_template_output_path, student_id, company_id)

                    print("Generating Endorsement Letter...")
                    gen_endorsement_letter(ojtrequirements_endorsement_letter, ojtrequirements_template_output_path, student_id, company_id, duration, **params)

                    print("Generating Medical Clearance...")
                    gen_medical(ojtrequirements_medical, ojtrequirements_template_output_path, student_id)

                    print("Generating MOA...")
                    gen_moa(ojtrequirements_moa, ojtrequirements_template_output_path, student_id=student_id, duration=duration, **params)

                    print("✅ All document generation functions called.")

                # ... (your existing except blocks for FileNotFoundError, KeyError, etc.) ...
                except Exception as doc_err: 
                    generation_successful = False
                    error_message = f"An unexpected error occurred during document generation: {doc_err}"
                    print(f"❌ {error_message}")
                    messages.error(request, f"Student assigned, but failed to generate documents: {doc_err}")

                # --- Check if files were created, then ZIP and REDIRECT ---
                if generation_successful:
                    try:
                        files_in_output = os.listdir(ojtrequirements_template_output_path)
                        generated_docs = [f for f in files_in_output if f.endswith('.docx') and student.lastname in f]

                        if not generated_docs:
                            warning_message = f"Student assigned, but no documents were found in '{ojtrequirements_template_output_path}'. Zip file will be empty."
                            print(f"⚠️ {warning_message}")
                            messages.warning(request, warning_message)
                            return redirect('ojt_hiring')
                        
                        # --- 1. ZIP THE FILES ---
                        zip_file_name = f"{student.studID}_{student.lastname}_OJT_Documents.zip"
                        
                        # This dynamically joins your MEDIA_ROOT with 'ojt_files'
                        zip_dir = os.path.join(settings.MEDIA_ROOT, 'ojt_files')
                        os.makedirs(zip_dir, exist_ok=True)
                        
                        zip_output_path = os.path.join(zip_dir, zip_file_name)
                        
                        with zipfile.ZipFile(zip_output_path, 'w') as zipf:
                            for doc in generated_docs:
                                file_path = os.path.join(ojtrequirements_template_output_path, doc)
                                zipf.write(file_path, arcname=doc)
                        
                        print(f"✅ Successfully created zip file: {zip_file_name}")

                        # --- 2. STORE INFO IN SESSION ---
                        request.session['download_url'] = f"{settings.MEDIA_URL}ojt_files/{zip_file_name}"
                        request.session['assigned_student_name'] = f"{student.firstname} {student.lastname}"
                        request.session['assigned_company_name'] = company.company_name
                        
                        # --- 3. REDIRECT TO THE NEW DOWNLOAD PAGE ---
                        return redirect('ojt_requiremets_download') 

                    except Exception as zip_err:
                        error_message = f"Student assigned, but failed to create zip file: {zip_err}"
                        print(f"❌ {error_message}")
                        messages.error(request, error_message)

                # If generation failed or no files found, redirect back
                return redirect('ojt_hiring')

            else: # Form is invalid
                print("❌ OJT Assignment Form is invalid:")
                if 'studID' in new_ojt_form.errors and any('already exists' in e for e in new_ojt_form.errors['studID']):
                    student_name = f"{student.firstname} {student.lastname}"
                    messages.warning(request, f"Warning: Student {student_name} (ID: {student_id}) is already assigned to an OJT company.")
                else:
                    print(new_ojt_form.errors)
                    messages.error(request, f"Failed to assign student: Invalid form data. {new_ojt_form.errors.as_text()}")
                return redirect('ojt_hiring')
                
                context = {
                    'form': OjtHiringForm(),
                    'assign_form': new_ojt_form,
                    'ojt_list': OJTCompany.objects.all(),
                    'base_template': base_template,
                }
                return render(request, 'jobplacement/ojthiring.html', context)

        except Exception as assign_err: # Catch errors during OJTStudent creation/saving
            error_message = f"Failed to save OJT assignment: {assign_err}"
            print(f"❌ {error_message}")
            messages.error(request, error_message)

    # If GET request or if POST failed before redirecting
    return redirect('ojt_hiring')

def student_suggestions(request):
    """Return a JSON list of students matching the query (by ID or name)."""
    query = request.GET.get('q', '').strip()
    if query:
        students = studentInfo.objects.filter(
            models.Q(studID__icontains=query) |
            models.Q(firstname__icontains=query) |
            models.Q(lastname__icontains=query)
        )[:10]  # limit results
        data = [
            {"id": s.studID, "name": f"{s.firstname} {s.lastname} ({s.studID})"}
            for s in students
        ]
    else:
        data = []
    return JsonResponse(data, safe=False)
# GI COMMENT KAY DILI NA JOBPLACEMENT GA HANDLE SA STUDENT LOGIN/SIGN UP

# 
# def student_signup_view(request):
#     if not (request.user.role != 'student' or request.user.is_superuser):
#         messages.info(request, 'Must be a superuser to access page')
#         return redirect('admin_login')

#     page = 'student_signup'
#     if request.method == 'POST':
#         form = StudentSignUpForm(request.POST)
#         if form.is_valid():
#             user = form.save()
#             backend = get_user_backend(user)
#             return redirect('home') 
#         else:
#             messages.error(request, 'There was an error with your sign-up.')
#     else:
#         form = StudentSignUpForm()
#     return render(request, 'jobplacement/student_login.html', {'form': form, 'page': page})

# def student_login(request):
#     page = 'student_login'
#     form = EmailAuthenticationForm()

#     if request.method == 'POST':
#         form = EmailAuthenticationForm(request, data=request.POST)
#         if form.is_valid():
#             email = form.cleaned_data.get('username')
#             password = form.cleaned_data.get('password')
#             user = authenticate(request, username=email, password=password)
#             if user is not None and isinstance(user, studentInfo):
#                 login(request, user)
#                 return redirect('home')  # Student dashboard URL
#             else:
#                 messages.error(request, 'Invalid email or password')
#         else:
#             messages.error(request, 'Invalid email or password')
#     else:
#         form = EmailAuthenticationForm()
#     return render(request, 'jobplacement/student_login.html', {'form': form, 'page': page})

# Create your views here.


def get_user_backend(user):
    if isinstance(user, JobPlacementAdminUser):
        return 'jobplacement.auth_backends.AdminUserBackend'
    elif isinstance(user, studentInfo):
        return 'jobplacement.auth_backends.studentInfoBackend'
    return None

# end of deletables\

# def admin_student(request):
#     return render(request, 'jobplacement/admin_and_student.html', {})

# def admin_login(request):
#     page='admin_login'
#     form = AdminLoginForm()
#     if request.method == 'POST':
#         form = AdminLoginForm(request, data=request.POST)
#         if form.is_valid():
#             email = form.cleaned_data.get('username')
#             password = form.cleaned_data.get('password')
#             user = authenticate(request, username=email, password=password)
#             if user is not None and user.role != 'student':
#                 backend = get_user_backend(user)
#                 login(request, user, backend=backend)
#                 return redirect('home')
#             else:
#                 messages.error(request, 'Invalid email or password')
#         else:
#             messages.error(request, 'Invalid email or password')
#     else:
#         form = AdminLoginForm()
#     return render(request, 'jobplacement/signin.html', {'page':page, 'form':form})


def logout_user(request):
    logout(request)
    return redirect('admin_student')

# (login_url='admin_student')

def mainpage(request):
    return redirect('ojt_hiring')

#   OJT HIRING THINGS
# (login_url='admin_student')

def ojt_hiring(request):
    
    ojt_hiring_form = OjtHiringForm()
    ojt_assign_form = OJTStudentForm()
    # This list is for the <datalist> dropdown, which needs ALL companies
    all_companies_list = OJTCompany.objects.all() 
    student_assignment = None
    download_url = None 

    base_template = "adminmain.html" if request.user.role != 'student' or request.user.is_superuser else "main.html"

    # --- Student Assignment Logic (No change) ---
    if request.user.is_authenticated and not (request.user.role != 'student' or request.user.is_superuser):
        try:
            student_info = studentInfo.objects.get(user=request.user) 
        except studentInfo.DoesNotExist:
            student_info = None
            
        if student_info:
            try:
                student_assignment = OJTStudent.objects.select_related('company_id').get(studID=student_info)
              
                zip_file_name = f"{student_info.studID}_{student_info.lastname}_OJT_Documents.zip"
                
                zip_file_path = os.path.join(settings.MEDIA_ROOT, 'ojt_files', zip_file_name)
                
                if os.path.exists(zip_file_path):
                    download_url = f"{settings.MEDIA_URL}ojt_files/{zip_file_name}"
                
            except OJTStudent.DoesNotExist:
                student_assignment = None

    # --- POST Logic (No change) ---
    if request.method == 'POST':
        if not (request.user.role != 'student' or request.user.is_superuser):
            messages.info(request, 'Must be staff/admin to access page')
            return redirect('admin_login')
        
        if not (isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser):
            messages.info(request, 'Must be Job Placement staff/admin to access page')
            return redirect('admin_login')
        
        form = OjtHiringForm(request.POST)
        if form.is_valid():
            newojt = form.save(commit=False)
            newojt.save()
            messages.success(request, "New OJT Hiring created successfully!")
            log_activity(request.user, "Created new OJT hiring")
            return redirect('ojt_hiring')  
        else:
            messages.error(request, "Form invalid")

    # --- GET Logic (Pagination, Search, Sort) ---
    search_query = request.GET.get('search', '')
    sort_by = request.GET.get('sort', 'date_desc') # Default sort is now Newest First

    # Start with all companies for filtering/sorting
    company_list = OJTCompany.objects.all()

    # Apply search filter
    if search_query:
        company_list = company_list.filter(
            Q(company_name__icontains=search_query) |
            Q(address__icontains=search_query) |
            Q(position__icontains=search_query) |
            Q(description__icontains=search_query)
        ).distinct()

    # Apply sorting
    if sort_by == 'date_asc':
        company_list = company_list.order_by('company_id') # Oldest first
    elif sort_by == 'name_asc':
        company_list = company_list.order_by('company_name')
    elif sort_by == 'name_desc':
        company_list = company_list.order_by('-company_name')
    elif sort_by == 'slots_asc':
        company_list = company_list.order_by('number_of_slots')
    elif sort_by == 'slots_desc':
        company_list = company_list.order_by('-number_of_slots')
    else: # Default is 'date_desc'
        company_list = company_list.order_by('-company_id') # Newest first

    # Apply pagination
    paginator = Paginator(company_list, 9) # 9 cards per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'form': ojt_hiring_form,
        'assign_form': ojt_assign_form,
        'ojt_list': all_companies_list, # Full list for the datalist
        'page_obj': page_obj,             # Paginated list for cards
        'student_assignment': student_assignment,
        'download_url': download_url,
        'base_template': base_template,
        'search_query': search_query,
        'sort_by': sort_by,
    }

    return render(request, 'jobplacement/ojthiring.html', context)

  

def ojthiring_delete(request, id):
    if not (request.user.role != 'student' or request.user.is_superuser):  # prevent student access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')
    
    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    

    try:
        company = OJTCompany.objects.get(company_id=id)
        company.delete()
    except Exception as e:
        messages.error(request, "Cannot delete company: There are students assigned to this company")
    
    return redirect('ojt_hiring')


def ojt_hiring_info(request, id):   # ojt company page
    if not (request.user.role != 'student' or request.user.is_superuser):  # prevent student access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')
    
    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    

    ojt = OJTCompany.objects.get(company_id = id)
    hiredStudents = OJTStudent.objects.filter(company_id__company_id=id)
    if request.method == 'POST':
        stud_id = request.POST.get('stud_id')

        student_obj = studentInfo.objects.get(studID = stud_id)
        existing_ojt = OJTStudent.objects.filter(student__studID=stud_id).last()
        if existing_ojt is None:
            form = OJTStudentForm(request.POST)
            if form.is_valid():
                newhired = form.save(commit=False)
                newhired.student = student_obj
                newhired.save()
                messages.success(request, "Student assigned")
            else:
                messages.error(request, "Failed to assign student")
        else:
            messages.info(request, "Student already under ojt")
    context = {'hired_students':hiredStudents, 'ojt':ojt}
    return render(request, 'jobplacement/ojthiring_info.html', context)

# (login_url='admin_student')

from ..models import OJTRequirements, studentInfo

from ..models import OJTRequirements, studentInfo

def ojtRequirements_tracker(request):
    # --- Base setup ---
    existing_requirement = None
    _reqform = OJTRequirementsForm()
    stat_widgets = StatusWidget()
    base_template = "adminmain.html" if (request.user.is_authenticated and (request.user.role != 'student' or request.user.is_superuser)) else "main.html"

    # --- Initialize Admin context ---
    req_records = []
    all_degrees = []
    search_query = ''
    selected_degree = ''
    sort_by = '' # <-- New variable for sorting

    if request.user.is_authenticated and (request.user.role != 'student' or request.user.is_superuser):
        # --- ADMIN VIEW LOGIC ---
        
        all_degrees = OJTRequirements.objects.select_related('student_id') \
                                              .order_by('student_id__degree') \
                                              .values_list('student_id__degree', flat=True) \
                                              .distinct()

        # Get filter parameters from the URL (using GET)
        search_query = request.GET.get('student_search', '')
        selected_degree = request.GET.get('degree_filter', '')
        sort_by = request.GET.get('sort_by', 'latest') # <-- Get sort_by, default to 'latest'

        # Start with all records
        req_records = OJTRequirements.objects.select_related('student_id')

        # Apply degree filter
        if selected_degree:
            req_records = req_records.filter(student_id__degree=selected_degree)

        # Apply search filter
        if search_query:
            req_records = req_records.annotate(
               full_name=Concat('student_id__firstname', Value(' '), 'student_id__lastname'),
               full_name_rev=Concat('student_id__lastname', Value(' '), 'student_id__firstname')
            )
            req_records = req_records.filter(
                Q(student_id__studID__icontains=search_query) |
                Q(full_name__icontains=search_query) |
                Q(full_name_rev__icontains=search_query)
            )
        
        # --- New Sorting Logic ---
        if sort_by == 'oldest':
            req_records = req_records.order_by('ojt_requirement_id') # Oldest first
        else:
            # Default to 'latest'
            req_records = req_records.order_by('-ojt_requirement_id') # Latest first
    
    elif request.user.is_authenticated:
        # --- STUDENT VIEW LOGIC ---
        try:
            student_info = studentInfo.objects.get(user=request.user)
            existing_requirement = OJTRequirements.objects.get(student_id=student_info)
            _reqform = OJTRequirementsForm(instance=existing_requirement)
            stat_widgets = StatusWidget(instance=existing_requirement)
        except (studentInfo.DoesNotExist, OJTRequirements.DoesNotExist):
            _reqform = OJTRequirementsForm()
            stat_widgets = StatusWidget()
            existing_requirement = None
    
    # --- Context for both Admin and Student ---
    context = {
        'form': _reqform,
        'existing_form': existing_requirement,
        'base_template': base_template,
        'status': stat_widgets,
        'req_records': req_records,
        
        # Admin filter context
        'all_degrees': all_degrees,
        'search_query': search_query,
        'selected_degree': selected_degree,
        'sort_by': sort_by, # <-- Pass sort_by to template
    }
    return render(request, 'jobplacement/ojt_requirements.html', context)

def ojt_requirements_submit(request):
    if request.method == 'POST':
        # --- Authentication Checks ---
        if not request.user.is_authenticated:
            messages.error(request, "You must be logged in as a student to submit requirements.")
            print("❌ User is not authenticated.")
            return redirect('ojt_requirements_tracker')

        try:
            student_info = studentInfo.objects.get(user=request.user)
        except studentInfo.DoesNotExist:
            messages.error(request, "No student profile found for your account.")
            print(f"❌ No studentInfo found for user: {request.user}")
            return redirect('ojt_requirements_tracker')

        # --- Form Handling ---
        existing_requirement = OJTRequirements.objects.filter(student_id=student_info).last()
        
        if existing_requirement:
            print(f"ℹ️ Updating existing requirement for {student_info}")
            form = OJTRequirementsForm(request.POST, request.FILES, instance=existing_requirement)
        else:
            print(f"ℹ️ Creating new requirement for {student_info}")
            form = OJTRequirementsForm(request.POST, request.FILES)

        # --- Validation and Saving ---
        if form.is_valid():
            print("✅ Form is valid. Proceeding to save.")
            requirement = form.save(commit=False)
            requirement.student_id = student_info
            
            # TRACKING MESSAGE FLAGS
            is_new_submission = not existing_requirement
            re_submitted_files = []

            # --- STATUS UPDATE LOGIC (from your original file) ---
            if existing_requirement:
                # Iterate over newly uploaded files
                for file_field_name in request.FILES.keys():
                    # Convert 'non_disclosure' -> 'nondis'
                    if file_field_name.endswith('_disclosure'): # Handles non_disclosure
                        attr_name = 'nondis'
                    elif file_field_name == 'parents_consent':
                        attr_name = 'consent'
                    elif file_field_name == 'application_letter':
                        attr_name = 'apl_letter'
                    else:
                        # For 'biodata', 'medical', 'moa', 'endorsement', 'certification'
                        attr_name = file_field_name.split('_')[0]
                    
                    # Map the short name to the status field name (e.g., 'nondis' -> 'nondis_stat')
                    status_field_name = f"{attr_name}_stat"

                    # Check current status in the database
                    current_db_status = getattr(existing_requirement, status_field_name, None)
                    
                    # If the file was previously DECLINED, set it back to PENDING
                    if current_db_status == OJTRequirements.DECLINED:
                        setattr(requirement, status_field_name, OJTRequirements.PENDING)
                        re_submitted_files.append(file_field_name.replace('_', ' ').title())

            requirement.save()
            log_activity(request.user, "Submitted OJT requirements") # Assumes log_activity is defined

            # --- MESSAGING (from your original file) ---
            if is_new_submission:
                messages.success(request, "All requirements submitted successfully! Awaiting review.")
            elif re_submitted_files:
                # Use warning for re-submission notification
                file_list = ", ".join(re_submitted_files)
                messages.warning(request, f"Files re-submitted: {file_list}. Awaiting administrator review.")
            elif request.FILES:
                messages.success(request, "New files successfully uploaded and sent for review.")
            else:
                messages.success(request, "Form data saved successfully.")

            print(f"✅ Requirement saved successfully for student {student_info}")
            return redirect('ojt_requirements_tracker')
        
        else:
            # --- THIS IS THE FIX ---
            # Re-render the page with the invalid form to show errors
            messages.warning(request, f"Only submit Image and pdf files.")
            
            print("❌ Form submission failed. Re-rendering page with errors.")
            print("📝 POST Data:", request.POST)
            print("📂 FILES Data:", request.FILES)
            # This will print the validation errors (e.g., "Invalid file type") to your console
            print("⚠️ Form Errors:", form.errors) 

            # This data is needed to render the template correctly
            existing_form = existing_requirement
            status_widgets = StatusWidget(instance=existing_form)
            base_template = "main.html" # Set base template for student view

            context = {
                'form': form,  # Pass the INVALID form back to the template
                'existing_form': existing_form,
                'status': status_widgets,
                'base_template': base_template,
            }
            # Re-render the page so the student sees the form errors
            return render(request, 'jobplacement/ojt_requirements.html', context)

    # If GET request, just redirect
    print("⚠️ GET request received, only POST is allowed here.")
    return redirect('ojt_requirements_tracker')
def ojt_requirements_accept(request):
    if not (request.user.role != 'student' or request.user.is_superuser): # allow access to admin and superuser only
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')

    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    

    if request.method == "POST":
        req_id = request.POST.get('req_id')
        attr_name = request.POST.get('attr_name')
        action = request.POST.get('action')
        try:
            req = OJTRequirements.objects.get(ojt_requirement_id = req_id)
            student = req.student_id
            if action == 'accept':              # Accept Action
                if attr_name == 'nondis':
                    try:
                        req.nondis_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: NON-DISCLOSURE AGREEMENT of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")

                elif attr_name == 'biodata':
                    try:
                        req.biodata_stat = OJTRequirements.ACCEPTED
                        req.save()
                        messages.success(request, "File Accepted")
                        log_activity(request.user, f"Accepted OJT Requirement: BIODATA of  {student.firstname} {student.lastname}")
                    except:
                        messages.error(request, "Action Failed")

                elif attr_name == 'consent':
                    try:
                        req.consent_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: PARENTS CONSENT of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")

                elif attr_name == 'medical':
                    try:
                        req.medical_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: MEDICAL of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")

                elif attr_name == 'apl_letter':
                    try:
                        req.apl_letter_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: APPLICATION LETTER of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")

                elif attr_name == 'moa':
                    try:
                        req.moa_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: MOA of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")

                elif attr_name == 'endorsement':
                    try:
                        req.endorsement_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: ENDORSEMENT LETTER of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")
                
                elif attr_name == 'cert':
                    try:
                        req.cert_stat = OJTRequirements.ACCEPTED
                        req.save()
                        log_activity(request.user, f"Accepted OJT Requirement: CERTIFICATION of  {student.firstname} {student.lastname}")
                        messages.success(request, "File Accepted")
                    except:
                        messages.error(request, "Action Failed")
            elif action == 'declined':             # Delete Action
                if attr_name == 'nondis':
                    req.nondis_stat = OJTRequirements.DECLINED
                    req.non_disclosure.delete()    # delete file and reference in database
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: NON-DISCLOSURE AGREEMENT of  {student.firstname} {student.lastname}")
                
                elif attr_name == 'biodata':
                    req.biodata_stat = OJTRequirements.DECLINED
                    req.biodata.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: BIODATA of  {student.firstname} {student.lastname}")
                
                elif attr_name == 'consent':
                    req.consent_stat = OJTRequirements.DECLINED
                    req.parents_consent.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: PARENTS CONSENT of  {student.firstname} {student.lastname}")
                
                elif attr_name == 'medical':
                    req.medical_stat = OJTRequirements.DECLINED
                    req.medical.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: MEDICAL of  {student.firstname} {student.lastname}")
                
                elif attr_name == 'apl_letter':
                    req.apl_letter_stat = OJTRequirements.DECLINED
                    req.application_letter.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: APPLICATION LETTER of  {student.firstname} {student.lastname}")
                
                elif attr_name == 'moa':
                    req.moa_stat = OJTRequirements.DECLINED
                    req.moa.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: MOA of  {student.firstname} {student.lastname}")
                
                elif attr_name == 'endorsement':
                    req.endorsement_stat = OJTRequirements.DECLINED
                    req.endorsement.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Letter Requirement: Endorsement Letter of  {student.firstname} {student.lastname}")
                elif attr_name == 'cert':
                    req.cert_stat = OJTRequirements.DECLINED
                    req.certification.delete()
                    messages.success(request, "File Declined")
                    log_activity(request.user, f"Declined OJT Requirement: CERTIFICATION of  {student.firstname} {student.lastname}")
        except:
            messages.error(request, "Requirement instance does not exist")
            return redirect('ojt_requirements_tracker')

    return redirect('ojt_requirements_tracker')


# SEMINAR THINGS
# (login_url='admin_student')

def seminar(request):
    seminars = Seminar.objects.all()
    form = SeminarForm()

    # Determine the layout (admin or student)
    base_template = "adminmain.html" if request.user.role != 'student' or request.user.is_superuser else "main.html"

    if request.method == 'POST':
        if not (request.user.role != 'student' or request.user.is_superuser):
            messages.info(request, 'Must be staff/admin to access page')
            return redirect('admin_login')

        if not (isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser):
            messages.info(request, 'Must be Jobplacement staff/admin to access page')
            return redirect('admin_login')

        form = SeminarForm(request.POST, request.FILES)
        if form.is_valid():
            newseminar = form.save(commit=False)
            newseminar.save()
            log_activity(request.user, f"Scheduled new seminar: {newseminar.title}")
            return redirect('seminar_page')

    context = {
        'seminars': seminars,
        'form': form,
        'base_template': base_template,
    }
    return render(request, 'jobplacement/seminar_page.html', context)


# 

def seminar_delete(request, id):
    if not (request.user.role != 'student' or request.user.is_superuser):  # prevent student access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')

    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')  
          
    seminar = Seminar.objects.get(seminar_id=id)
    seminar.delete()
    return redirect('seminar_page')

# 

def manage_attendance(request, id): # seminar attendance page
    if not (request.user.role != 'student' or request.user.is_superuser): # prevent student/alien access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')

    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    
    
    sem_form = SeminarAttendanceForm()
    attendance = SeminarAttendance.objects.filter(seminar_id__seminar_id=id, ispending=True, attended=False)
    present_students = SeminarAttendance.objects.filter(attended=True, seminar_id__seminar_id = id)
    
    print(attendance)
    context = {'attendance':attendance, 'sem_id':id, 'sem_form':sem_form, 'presents':present_students}
    return render(request, 'jobplacement/sem_att_manager.html', context)
# ...existing code...

def pend_attendance(request):

    # handles student attendance request
    if request.method == 'POST':
        try:
          print("pend_attendance POST keys:", dict(request.POST))
          print("pend_attendance COOKIES:", request.COOKIES.keys())
        except Exception as _:
          print("pend_attendance: failed to dump POST data")

        stud_id = request.POST.get('student_id')
        sem_id = request.POST.get('seminar_id')

        if not stud_id or not sem_id:
            messages.error(request, "Student ID and Seminar ID are required.")
            return redirect('manage_att', sem_id)

        try:
            # Try to find existing attendance using the FK field name 'student_id'
            sem_att = SeminarAttendance.objects.get(
                student_id__studID=stud_id,
                seminar_id__seminar_id=sem_id
            )
            sem_att.ispending = True
            sem_att.save()
            messages.success(request, "Attendance updated successfully.")
            log_activity(request.user, "Attendance Updated")
            return redirect('manage_att', id=sem_id)
        except SeminarAttendance.DoesNotExist:
            # If the attendance record does not exist, create a new one
            try:
                student_obj = studentInfo.objects.get(studID=stud_id)
            except studentInfo.DoesNotExist:
                messages.error(request, "Student does not exist.")
                return redirect('manage_att', id=sem_id)
            try:
                seminar_obj = Seminar.objects.get(seminar_id=sem_id)
            except Seminar.DoesNotExist:
                messages.error(request, "Seminar does not exist.")
                return redirect('manage_att', id=sem_id)

            try:
                # Use the correct FK field names (student_id, seminar_id) and avoid 'studID'
                att, created = SeminarAttendance.objects.get_or_create(
                    student_id=student_obj,
                    seminar_id=seminar_obj,
                    defaults={'ispending': True, 'attended': False}
                )
                if not created:
                    att.ispending = True
                    att.save()
                print(f"pend_attendance: attendance record created={created} id={att.pk}")
                # verify immediately (debug)
                qs = SeminarAttendance.objects.filter(seminar_id__seminar_id=sem_id, ispending=True, attended=False)
                print("pend_attendance: pending queryset count:", qs.count(), qs[:5])
                log_activity(request.user, "New attendance recorded" if created else "Attendance Updated")
                messages.success(request, "New attendance recorded successfully." if created else "Attendance updated successfully.")
                return redirect('manage_att', id=sem_id)
            except Exception as e:
                print("pend_attendance: error creating attendance:", e)
                messages.error(request, f"Error creating new attendance: {e}")
                return redirect('manage_att', id=sem_id)

    return redirect('seminar')
# ...existing code...

# 

def cancel_pending(request, id): # handles attendance request cancel
    if not (request.user.role != 'student' or request.user.is_superuser):
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')    
    
    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')        
    
    if request.method == 'POST':
        try:
            att = SeminarAttendance.objects.get(sem_at_id = id)
            att.ispending = False
            att.save()
            return redirect('manage_att', id=att.seminar_id.seminar_id)
        except ObjException:
            messages.error(request, "Object does not exist")
        return redirect('home')

@csrf_exempt

def attend_all_pending(request, id): # handle attend all button
    if not (request.user.role != 'student' or request.user.is_superuser): # prevent student/alien access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')
     
    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            print(f"Datas: {data}")
            for item in data:
                print(f"Item: {item}")
                sem_id = item.get('sem_att_id')
                print(f"sem_id: {sem_id}")
                att_li = SeminarAttendance.objects.get(sem_at_id = sem_id, ispending=True, attended=False)
                att_li.ispending = False
                att_li.attended = True
                att_li.save()
                
        except SeminarAttendance.DoesNotExist:
            print("Seminar Attendance does not exist")
            return JsonResponse({"error": "Seminar Attendance id does not exist"})
        except Exception as e:
            print(e)
            return JsonResponse({"error": "Failed to mark student attendance"})

        return JsonResponse({"success": "Attendance all successful"})
    return redirect('manage_att', id=id)

    # NON-ACADEMIC AWARD ISSUANCE


def non_acad_page(request):
    if not (request.user.role != 'student' or request.user.is_superuser):
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')
    
    students_list = studentInfo.objects.all()
     
    
    def fill_placeholders(doc, data): # replace placeholders from template
        for p in doc.paragraphs:
            for key, value in data.items():
                if key in p.text:
                    inline = p.runs
                    for item in inline:
                        if key in item.text:
                            item.text = item.text.replace(key, value)

    def generate_document(template_path, output_path, data): # generate document for student
        try:
            doc = Document(template_path)
            fill_placeholders(doc, data)
            doc.save(output_path)
            # messages.success(request, "Successfully generated Docs")
        except:
            messages.error(request, "Failed to generate Docs")

    # Processes Non-Academic Award
    if request.method == 'POST':
        student_id = request.POST.get('student_id')
        date_issued_str = request.POST.get('date_issued') # <-- MOVED HERE
        
        try:
            # --- TRY ONLY THE CODE THAT CAN FAIL ---
            student = studentInfo.objects.get(studID=student_id)
            
        except studentInfo.DoesNotExist: 
            messages.error(request, f"Student with ID '{student_id}' not found. Please select a valid student.")
            return redirect('non_acad_page') # <-- ADD THIS to stop the function
        except Exception as e:
            # Catch other errors (like if student_id is None)
            messages.error(request, f"Failed to get inputs: {e}")
            
            
        # Default value if date_issued_str is None or empty
        default_date = datetime.now()   

        # Parse the date string if it's not None or empty, else use the default date
        if date_issued_str:
            date_issued = datetime.strptime(date_issued_str, '%Y-%m-%d').date()
        else:
            date_issued = default_date        

        formatted_date_issued = date_issued.strftime('%B %d, %Y')
        award = request.POST.get('award', '').strip()
        program = request.POST.get('program', '').strip().upper()

        achievement = request.POST.get('achievement')
        try:
            print('generating data')
            data = {
                '[date_given]': formatted_date_issued,
                '[firstname]': student.firstname,
                '[lastname]': student.lastname,

                # OJT
                '[position]': request.POST.get('position'),
                '[company]': request.POST.get('company_name'),
                '[company_address]': request.POST.get('company_address'),

                # CAPSTONE
                '[capstone_title]': request.POST.get('capstone_title'),

                # Leadership
                '[academic_year]': request.POST.get('acad_year'),

                # Research Fields
                '[research_title]': request.POST.get('research_title'),

                # Others
                '[achievement]': f"{achievement.upper()}",

            }
        except:
            messages.error(request, "Error in datas")

        try: 
            print(f'{program}')
            print(f'{award}')
            print(f"COED:{program == 'COED'}, {award == 'Researcher of the Year'}")

            if program == 'AGRICULTURE' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/AGRICULTURE/leadership_award_AGRICULTURE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_Leadership_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_Leadership_Award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'AGRICULTURE' and award == 'Social Responsibility and Civic Engangement Award':
                
                template_path = './media/templates/non_academic_awards/AGRICULTURE/social_responsibility_and_civic_engangement_award_AGRICULTURE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_Social_Responsibility_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_Social_Responsibility_Award.docx'
                    log_activity(user=request.user, action=f"Issued Social Responsibility and Civic Engangement Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'AGRICULTURE' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/AGRICULTURE/new_achievement_AGRICULTURE.docx'
                output_path = f"./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx"
                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued Social {achievement} Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'BIT' and award == 'Best OJT Award':
                template_path = './media/templates/non_academic_awards/BIT/best_OJT_awards_BIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_best_OJT_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_best_OJT_award.docx'
                    log_activity(user=request.user, action=f"Issued Best OJT Award to {student.firstname} {student.lastname}")
                    return response
            
            elif program == 'BIT' and award == 'Researcher of the Year':
                template_path = './media/templates/non_academic_awards/BIT/researcher_of_the_year_award_BIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_researcher_of_the_year_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_researcher_of_the_year_award.docx'
                    log_activity(user=request.user, action=f"Issued Researcher of the year Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BIT' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/BIT/new_achievement_BIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BES' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/BES/leadership_award_ENVIRONMENTAL_SCIENCE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'BES' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/BES/new_achievement_BES.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'BSHM' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/BSHM/leadership_award_BSHM.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'BSHM' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/BSHM/new_achievement_BSHM.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'BSIE' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/BSIE/leadership_award_BSIE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIE' and award == 'Outstanding Athlete Award':
                template_path = './media/templates/non_academic_awards/BSIE/outstanding_athlete_award_BSIE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_outstanding_athlete_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_outstanding_athlete_award.docx'
                    log_activity(user=request.user, action=f"Issued Outstanding Athletic Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIE' and award == 'Researcher of the Year':
                template_path = './media/templates/non_academic_awards/BSIE/researcher_of_the_year_award_BSIE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_researcher_of_the_year_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_researcher_of_the_year_award.docx'
                    log_activity(user=request.user, action=f"Issued Researcher of the year Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIE' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/BSIE/new_achievement_BSIE.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIT' and award == 'Best Capstone':
                template_path = './media/templates/non_academic_awards/BSIT/best_capstone_award_BSIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_best_capstone_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_best_capstone_award.docx'
                    log_activity(user=request.user, action=f"Issued Best Capstone Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIT' and award == 'Excellence Award':
                template_path = './media/templates/non_academic_awards/BSIT/excellence_award_BSIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_excellence_award_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_excellence_award_award.docx'
                    log_activity(user=request.user, action=f"Issued Excellence Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIT' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/BSIT/leadership_award_BSIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'BSIT' and award == 'Programmer of the Year':
                template_path = './media/templates/non_academic_awards/BSIT/programmer_of_the_year_award_BSIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_programmer_of_the_year_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_programmer_of_the_year_award.docx'
                    log_activity(user=request.user, action=f"Issued Programmer of the Year Award to {student.firstname} {student.lastname}")
                    return response
            
            elif program == 'BSIT' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/BSIT/new_achievement_BSIT.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
            
            elif program == 'CAS' and award == 'Academic Leadership Award':
                template_path = './media/templates/non_academic_awards/CAS/academic_leadership_award_CAS.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_academic_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_academic_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Academic Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'CAS' and award == 'BAEL Pride Award':
                template_path = './media/templates/non_academic_awards/CAS/bael_pride_award_CAS.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_bael_pride_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_bael_pride_award.docx'
                    log_activity(user=request.user, action=f"Issued Bael Pride Award to {student.firstname} {student.lastname}")
                    return response
                                
            elif program == 'CAS' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/CAS/leadership_award_CAS.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'CAS' and award == 'Loyalty Award':
                template_path = './media/templates/non_academic_awards/CAS/loyalty_award_CAS.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_loyalty_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_loyalty_award.docx'
                    log_activity(user=request.user, action=f"Issued Loyalty Award to {student.firstname} {student.lastname}")
                    return response
            
            elif program == 'CAS' and award == 'Outstanding Athlete Award':
                template_path = './media/templates/non_academic_awards/CAS/outstanding_athlete_award_CAS.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_outstanding_athlete_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_outstanding_athlete.docx'
                    log_activity(user=request.user, action=f"Issued Outstanding Athlete Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'CAS' and award == 'Others':
                
                template_path = './media/templates/non_academic_awards/CAS/new_achievement_CAS.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
        
            elif program == 'COED' and award == 'Best in Elocution Award':
                template_path = './media/templates/non_academic_awards/COED/best_in_elocution_award_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_best_in_elocution_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_best_in_elocution.docx'
                    log_activity(user=request.user, action=f"Issued Best in Elocution Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'COED' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/COED/leadership_award_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'COED' and award == 'Relentless Mentor of the Year Award':
                template_path = './media/templates/non_academic_awards/COED/relentless_mentor_of_the_year_award_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_relentless_mentor_of_the_year_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_relentless_mentor_of_the_year.docx'
                    log_activity(user=request.user, action=f"Issued Relentles Mentor of the Year Award to {student.firstname} {student.lastname}")
                    return response
            
            elif program == 'COED' and award == 'Researcher of the Year':
                template_path = './media/templates/non_academic_awards/COED/researcher_of_the_year_award_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_researcher_of_the_year_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_researcher_of_the_year.docx'
                    log_activity(user=request.user, action=f"Issued Researcher of the Year Award to {student.firstname} {student.lastname}")
                    return response

            elif program == 'COED' and award == 'Student Extensionista of the Year Award':
                template_path = './media/templates/non_academic_awards/COED/student_extensionista_of_the_year_award_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_student_extensionista_of_the_year_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_student_extensionista_of_the_year.docx'
                    log_activity(user=request.user, action=f"Issued Student Extensionista of the Year Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'COED' and award == 'Others':
                template_path = './media/templates/non_academic_awards/COED/new_achievement_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'FORESTRY' and award == 'Leadership Award':
                template_path = './media/templates/non_academic_awards/COED/leadership_award_COED.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_leadership_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_leadership_award.docx'
                    log_activity(user=request.user, action=f"Issued Leadership Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'FORESTRY' and award == 'Outstanding Athlete':
                template_path = './media/templates/non_academic_awards/FORESTRY/outstanding_athlete_award_FORESTRY.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_outstanding_athlete_award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_outstanding_athlete_award.docx'
                    log_activity(user=request.user, action=f"Issued Outstanding Athlete Award to {student.firstname} {student.lastname}")
                    return response
                
            elif program == 'FORESTRY' and award == 'Others':
                template_path = './media/templates/non_academic_awards/FORESTRY/new_achievement_FORESTRY.docx'
                output_path = f'./media/templates/non_academic_awards/output/{student.firstname}_{student.lastname}_{achievement}_Award.docx'

                generate_document(template_path, output_path, data)
                print('passed generate document')
                # Optionally return a response with the document, like a download link
                with open(output_path, 'rb') as doc_file:
                    response = HttpResponse(doc_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    response['Content-Disposition'] = f'attachment; filename={student.firstname}_{student.lastname}_{achievement}_Award.docx'
                    log_activity(user=request.user, action=f"Issued {achievement} Award to {student.firstname} {student.lastname}")
                    return response
        except:
            messages.error(request, "Failed to create achievement")
    context = {'students': students_list,}
    return render(request, 'jobplacement/non_academic.html', context)

    # TRANSACTION REPORTS THINGS

def log_activity(user, action): # transaction report auto record
    user_type = 'admin' if user.role != 'student' else 'student'
    TransactionReport.objects.create(
        user=user,
        action=action,
        date_created=timezone.now(),
        user_type=user_type
    )

# 

def transRep(request): # transaction report page
    if not (request.user.role != 'student' or request.user.is_superuser): # prevent student/alien access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')    
    
    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access

        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    

    form = TransactionForm()
    records = TransactionReport.objects.all()
    dt = datetime.now().strftime('%Y-%m-%dT%H:%M')
    print(f"Date time now: {dt}")
    _monthly = False

    # handle filter function
    if request.method == 'POST':
        monthly_filter = request.POST.get('monthly_filter')
        date_time = request.POST.get('date_time_filter')
        day = month = year = None

        try: 
            date_time = datetime.strptime(date_time, '%Y-%m-%dT%H:%M')
            dt = date_time.strftime('%Y-%m-%dT%H:%M')

            try:
                day = date_time.day
                month = date_time.month
                year = date_time.year
            except:
                messages.error(request, f"Invalid: {day}, {month}, {year}")
        except ValueError:
            messages.error(request, "Invalid date time")
        print(monthly_filter)
        if monthly_filter == 'true':
            print("monthly filter")
            _monthly = True
            records = records.filter(date_created__month = month, date_created__year = year)
        else:
            print("daily filter")
            _monthly = False
            records = records.filter(date_created__month = month, date_created__year=year, date_created__day=day)
    
    context = {'transactions':records, 'date_time':dt,'prev_period':_monthly, 'form':form}
    return render(request, 'jobplacement/trans_report.html', context )


def transRep_print(request): 
    if not (request.user.role != 'student' or request.user.is_superuser): # prevent student/alien access
        messages.info(request, 'Must be staff/admin to access page')
        return redirect('admin_login')
    
    if not ( isinstance(request.user, JobPlacementAdminUser) or request.user.is_superuser): # prevent other admin access
        messages.info(request, 'Must be Jobplacement staff/admin to access page')
        return redirect('admin_login')    

    records = TransactionReport.objects.all()
    dt = datetime.now().strftime('%B %d, %Y %H:%M:%S')
    filter_type = "None"

    # Print transaction report
    if request.method == 'POST':
        monthly_filter = request.POST.get('period_filter')
        date_time = request.POST.get('date_filter')

        day = month = year = None
        try: 
            date_time = datetime.strptime(date_time, '%Y-%m-%dT%H:%M')

            try:
                day = date_time.day
                month = date_time.month
                year = date_time.year
            except:
                messages.error(request, f"Invalid: {day}, {month}, {year}")
        except ValueError:
            messages.error(request, "Invalid date time")
        
        if monthly_filter == 'True':
            filter_type = "MONTHLY"
            print("monthly filter")
            records = records.filter(date_created__month = month, date_created__year = year)
        else:
            filter_type = "DAILY"
            print("daily filter")
            records = records.filter(date_created__month = month, date_created__year=year, date_created__day=day)

    context = {'transactions':records, 'datetime':dt, 'filter_type':filter_type}
    # return render(request, 'jobplacement/ready_to_print2.html', context)

    html_string = render_to_string('jobplacement/ready_to_print2.html', context)

    base_url = request.build_absolute_uri(os.path.join(os.path.dirname(request.path), 'static'))
    
    pdf_file = HTML(string=html_string, base_url=base_url).write_pdf()
    
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment;filename="TransactionReport.pdf"'

    return response 

# suggest students with matching query
class search_suggestions(View):
    def get(self, request):
        query = request.GET.get('query', '') # query comes from ajax
        suggestions = []
        if query:
            suggestions = list(studentInfo.objects.filter(
                Q(studID__icontains=query) |
                Q(firstname__icontains = query) |
                Q(lastname__icontains = query)
            ).values_list('studID', 'lastname', 'firstname'))
            return JsonResponse(suggestions, safe=False)

# suggest companies with matching query
class company_suggestions(View):
    print('company suggest')
    def get(self, request):
        query = request.GET.get('query', '') # query comes from ajax
        suggestions = []
        if query:
            suggestions = list(OJTCompany.objects.filter(
                Q(company_id__icontains=query) |
                Q(company_name__icontains = query) |
                Q(owner__icontains = query)
            ).values_list('company_id', 'company_name', 'owner'))
            return JsonResponse(suggestions, safe=False)



def del_ojt(request):
    if OJTRequirements.objects.all().delete():
        messages.success(request, "Clear success")
    else:
        messages.error(request, "Clear failed")
    
    return redirect('home')


# generate application letter document

def gen_application_letter(template_path, output_path, student_id, company_id):
    student = studentInfo.objects.get(studID = student_id)
    company = OJTCompany.objects.get(company_id = company_id)
    current_date = datetime.now().strftime("%B %d, %Y")

    th = ""
    if student.yearlvl == 1:
        th = "st"
    elif student.yearlvl == 2:
        th = "nd"
    elif student.yearlvl == 3:
        th = "rd"
    else:
        th = 'th'

    data = {
        '[date]': f"{current_date}",
        '[position]': f"{company.position}",
        '[name_of_company_representative]': f"{company.owner}",
        '[company_name]': f"{company.company_name}",
        '[company_address]': f"{company.address}",
        '[degree_and_year_of_student]': f"{student.degree} as {student.yearlvl}{th} student",
        '[contact_number_of_student]': f"{student.contact}",
        '[firstname]': f"{student.firstname}",
        '[lastname]': f"{student.lastname}",
    }
    try:
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
    except:
        print("Failed to generate letter")

    try:
        new_path = os.path.join(output_path, f"{student.lastname}_ApplicationLetter_{current_date}.docx")

        doc.save(new_path)
    except:
        print("failed to save file")

    print("Application Letter Generated")

# generate biodata document

def gen_biodata(template_path, output_path, student_id, company_id):
    print("generating biodata")
    student = studentInfo.objects.get(studID = student_id)
    company = OJTCompany.objects.get(company_id = company_id)
    print(student)
    print(company)
    current_date = datetime.now().strftime("%B %d, %Y")

    try:
        doc = Document(template_path)
        new_path = os.path.join(output_path, f"{student.lastname}_biodata_{current_date}.docx")

        doc.save(new_path)
    except:
        print("failed to save file")

    print("Biodata Generated")

# generate endorsement letter document

def gen_endorsement_letter(template_path, output_path, student_id, company_id, duration, **params):
    print("generating endorsement")
    student = studentInfo.objects.get(studID = student_id)
    company = OJTCompany.objects.get(company_id = company_id)
    current_date = datetime.now().strftime("%B %d, %Y")

    data = {
        '[date]': current_date,
        '[name_of_company_representative]': company.owner,
        '[position]': company.position,
        '[company_name]': company.company_name,
        '[company_address]': company.address,
        '[hours_needed]': duration,
        '[firstname]': student.firstname,
        '[lastname]': student.lastname,
        '[endorser_phonenum]': params.get('endorser_num'),
        '[endorser_email]': params.get('endorser_email'),
        '[endorser_name]': params.get('endorser_name'),
        '[program]': params.get('endorser_program'),
    }
    try:
        print('saving letter')
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
    except:
        print("Failed to generate letter")

    try:
        new_path = os.path.join(output_path, f"{student.lastname}_Endorsement{current_date}.docx")

        doc.save(new_path)
    except:
        print("failed to save file")

    print("Endorsement Generated")

# generate medical document

def gen_medical(template_path, output_path, student_id):
    print('Generating Medical')
    student = studentInfo.objects.get(studID = student_id)
    print(student)
    current_date = datetime.now().strftime("%B %d, %Y")

    try:
        doc = Document(template_path)
        new_path = os.path.join(output_path, f"{student.lastname}_medical_{current_date}.docx")

        doc.save(new_path)
    except:
        print("failed to save file")

    print("Medical Generated")

# generate moa document

def gen_moa(template_path, output_path, student_id, duration, **params):
    print('generating MOA')
    student = studentInfo.objects.get(studID = student_id)
    current_date = datetime.now()
    day = current_date.strftime("%d")
    month = current_date.strftime("%B")
    year = current_date.strftime("%Y")
    date_for_filename = current_date.strftime("%Y%m%d")
    data = {
        '[date]': day,
        '[month]': month,
        '[year]': year,
        '[name_of_coordinator]': params.get('endorser_name'),
        '[firstname]': student.firstname,
        '[lastname]': student.lastname,
        '[degree_program]': student.degree,
        '[hours]': duration,
    }
    doc = Document(template_path)
    try:
        print('saving letter')
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                        print(f"{paragraph.text}: {value}")

        print('updating table')
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
    except:
        print("Failed to generate letter")

    try:
        new_path = os.path.join(output_path, f"{student.lastname}_MOA_{date_for_filename}.docx")

        doc.save(new_path)
    except:
        print("failed to save file")

    print("MOA Generated")

# zip folder

def zip_files_in_folder(folder_path):
    print("buffing zip folder")
    zip_buffer = BytesIO()
    
    with ZipFile(zip_buffer, 'w') as zip_file:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zip_file.write(file_path, os.path.relpath(file_path, folder_path))
    
    zip_buffer.seek(0)
    return zip_buffer

# download zip folders

def download_zipped_folder(request, folder_path):
    print("downlodaing zip folder")
    try:
        zip_buffer = zip_files_in_folder(folder_path)
        zip_filename = os.path.basename(folder_path.rstrip('/')) + '.zip'
        
        response = FileResponse(zip_buffer, as_attachment=True, filename=zip_filename)
        
        # Delete files after downloading
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                os.remove(os.path.join(root, file))

        return response
    
    except Exception as e:
        print(f"Failed to zip and download folder: {e}")
        raise Http404("Failed to generate zip file.")

# download zipped folder

def download_zipped_ojt_templates(request):
    print("downloading ojt_templates")
    folder_path = './media/templates/ojt_requirements/generated'
    return download_zipped_folder(request, folder_path)

# as the function name

def file_scrapper(request):
    if request.method == 'POST':
        file = ScrapperFile( request.POST, request.FILES)
        if file.is_valid():
            csv_file = request.FILES['file']
            decoded_file = csv_file.read().decode('utf-8').splitlines()
            reader = csv.DictReader(decoded_file)

            for row in reader:
                studentInfo.objects.create(
                    studID=row['studID'],
                    lrn = row['lrn'],
                    firstname=row['firstname'],
                    lastname= row['lastname'],
                    middlename=row['middlename'],
                    program=row['degree'],
                    yearlvl = row['yearlvl'],
                    sex = row['sex'],
                    email = row['emailadd'],
                    contact = row['contact']
                )

            messages.success(request, "File scrapper success")
            return redirect('home')
        
    else:
        form = ScrapperFile()
    return render(request, 'jobplacement/scrapper.html', {"form":form})